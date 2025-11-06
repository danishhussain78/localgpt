from django.shortcuts import render, redirect
from django.http import JsonResponse
from django.contrib import auth, messages
from django.contrib.auth.models import User
from django.utils import timezone
from django.core.files.storage import default_storage
from .models import Chat, Document

import requests
import json
import os
import PyPDF2
import docx
from sentence_transformers import SentenceTransformer
import chromadb
from chromadb.config import Settings
from dotenv import load_dotenv
from datetime import datetime

# ================================
# Load environment variables
# ================================
load_dotenv()

# Use environment variable if provided
OLLAMA_HOST = os.getenv("OLLAMA_HOST", "http://localhost:11434")

# Initialize embedding model
embedding_model = SentenceTransformer("all-MiniLM-L6-v2")

# Initialize ChromaDB client with persistent storage
chroma_client = chromadb.PersistentClient(path="./chroma_db")


# ================================
# Helper Functions
# ================================
def get_user_collection_name(user_id):
    """Generate unique collection name for each user."""
    return f"user_{user_id}_documents"


def get_unique_chunk_id(user_id, document_id, chunk_index):
    """Generate globally unique chunk ID."""
    return f"u{user_id}_doc{document_id}_chunk{chunk_index}"


# ================================
# Document Text Extraction
# ================================
def extract_text_from_pdf(file_path):
    """Extract text from PDF file."""
    text = ""
    try:
        with open(file_path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
    except Exception as e:
        print(f"Error extracting PDF: {e}")
    return text


def extract_text_from_docx(file_path):
    """Extract text from DOCX file."""
    text = ""
    try:
        doc = docx.Document(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
    except Exception as e:
        print(f"Error extracting DOCX: {e}")
    return text


def extract_text_from_txt(file_path):
    """Extract text from TXT file."""
    try:
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read()
    except UnicodeDecodeError:
        try:
            with open(file_path, "r", encoding="latin-1") as file:
                return file.read()
        except Exception as e:
            print(f"Error extracting TXT: {e}")
            return ""
    except Exception as e:
        print(f"Error extracting TXT: {e}")
        return ""


def chunk_text(text, chunk_size=500, overlap=50):
    """Split text into overlapping chunks."""
    if not text or not text.strip():
        return []
    
    words = text.split()
    if len(words) == 0:
        return []
    
    chunks = []
    for i in range(0, len(words), chunk_size - overlap):
        chunk = " ".join(words[i:i + chunk_size])
        if chunk.strip():
            chunks.append(chunk.strip())
    
    return chunks


# ================================
# Document Processing
# ================================
def process_document(document_id, user_id):
    """Process document and store in user-specific vector database."""
    try:
        document = Document.objects.get(id=document_id, user_id=user_id)
        file_path = document.file.path

        # Extract text based on file type
        if file_path.endswith(".pdf"):
            text = extract_text_from_pdf(file_path)
        elif file_path.endswith(".docx"):
            text = extract_text_from_docx(file_path)
        elif file_path.endswith(".txt"):
            text = extract_text_from_txt(file_path)
        else:
            print(f"Unsupported file type: {file_path}")
            return False

        if not text or not text.strip():
            print(f"No text extracted from document {document_id}")
            return False

        # Chunk the text
        chunks = chunk_text(text)
        
        if not chunks:
            print(f"No chunks created from document {document_id}")
            return False

        # Get or create user-specific collection
        collection_name = get_user_collection_name(user_id)
        
        try:
            collection = chroma_client.get_collection(name=collection_name)
        except Exception:
            collection = chroma_client.create_collection(
                name=collection_name,
                metadata={"user_id": str(user_id)}
            )

        # Store chunks with user-specific IDs
        embeddings_list = []
        documents_list = []
        ids_list = []
        metadatas_list = []

        for idx, chunk in enumerate(chunks):
            try:
                embedding = embedding_model.encode(chunk).tolist()
                chunk_id = get_unique_chunk_id(user_id, document_id, idx)
                
                embeddings_list.append(embedding)
                documents_list.append(chunk)
                ids_list.append(chunk_id)
                metadatas_list.append({
                    "user_id": str(user_id),
                    "document_id": str(document_id),
                    "document_title": document.title,
                    "chunk_index": idx,
                    "timestamp": datetime.now().isoformat()
                })
            except Exception as e:
                print(f"Error processing chunk {idx}: {e}")
                continue

        # Batch add all chunks
        if embeddings_list:
            try:
                collection.add(
                    embeddings=embeddings_list,
                    documents=documents_list,
                    ids=ids_list,
                    metadatas=metadatas_list
                )
            except Exception as e:
                print(f"Error adding to collection: {e}")
                return False

        # Mark document as processed
        document.is_processed = True
        document.save()
        
        print(f"Successfully processed document {document_id} with {len(chunks)} chunks")
        return True
        
    except Document.DoesNotExist:
        print(f"Document {document_id} not found for user {user_id}")
        return False
    except Exception as e:
        print(f"Error processing document: {e}")
        import traceback
        traceback.print_exc()
        return False


def search_documents(user_id, query, top_k=3):
    """Search for relevant chunks in user-specific documents only."""
    try:
        collection_name = get_user_collection_name(user_id)
        collection = chroma_client.get_collection(name=collection_name)
        
        # Generate query embedding
        query_embedding = embedding_model.encode(query).tolist()
        
        # Search with user filter
        results = collection.query(
            query_embeddings=[query_embedding],
            n_results=top_k,
            where={"user_id": str(user_id)}
        )

        if results and results["documents"] and len(results["documents"]) > 0:
            return results["documents"][0]
        return []
        
    except Exception as e:
        print(f"Error searching documents for user {user_id}: {e}")
        return []


def delete_user_document_chunks(user_id, document_id):
    """Delete all chunks for a specific document."""
    try:
        collection_name = get_user_collection_name(user_id)
        collection = chroma_client.get_collection(name=collection_name)
        
        # Get all chunks for this user
        results = collection.get(where={"user_id": str(user_id)})
        
        # Filter out only the chunks for this document
        ids_to_delete = [
            _id for _id, meta in zip(results.get("ids", []), results.get("metadatas", []))
            if meta.get("document_id") == str(document_id)
        ]
        
        if ids_to_delete:
            collection.delete(ids=ids_to_delete)
            print(f"Deleted {len(ids_to_delete)} chunks for document {document_id}")
            return True
        return False
        
    except Exception as e:
        print(f"Error deleting document chunks: {e}")
        return False



# ================================
# LLaMA / Ollama Chat Function (FIXED)
# ================================
def ask_ollama(message, context_chunks=None):
    """
    Sends a message to Ollama with optional document context.
    Uses the correct /api/generate endpoint.
    """
    # Build the prompt
    if context_chunks and len(context_chunks) > 0:
        context_text = "\n\n".join(
            [f"Context {i+1}:\n{chunk}" for i, chunk in enumerate(context_chunks)]
        )
        prompt = f"""You are a helpful assistant. Use the following context from the uploaded documents to answer the user's question. If the answer is not in the context, say so.

Document Context:
{context_text}

User Question: {message}

Answer:"""
    else:
        prompt = message

    # Ollama /api/generate endpoint format
    payload = {
        "model": "llama3.2:1b",
        "prompt": prompt,
        "stream": False,
        "options": {
            "temperature": 0.7,
            "top_p": 0.9
        }
    }

    try:
        response = requests.post(
            f"{OLLAMA_HOST}/", 
            json=payload, 
            timeout=60
        )
        response.raise_for_status()
        data = response.json()

        # Ollama returns "response" field
        if "response" in data:
            return data["response"]
        else:
            print(f"Unexpected response format: {data}")
            return "Unexpected response format from Ollama."

    except requests.exceptions.Timeout:
        return "⏱️ Request timed out. Please try again."
    except requests.exceptions.ConnectionError:
        return f"❌ Could not connect to Ollama at {OLLAMA_HOST}. Please make sure Ollama is running."
    except requests.exceptions.HTTPError as e:
        print(f"HTTP Error: {e}")
        print(f"Response content: {response.text if response else 'No response'}")
        return f"Error communicating with Ollama: {str(e)}"
    except Exception as e:
        print(f"Ollama API error: {e}")
        import traceback
        traceback.print_exc()
        return "Sorry, I couldn't process your request. Please check if Ollama is running."


# ================================
# Chatbot Views
# ================================
def chatbot(request):
    """Main chatbot page with user-specific data."""
    chats = Chat.objects.filter(user=request.user).order_by("created_at")
    documents = Document.objects.filter(user=request.user).order_by("-uploaded_at")

    if request.method == "POST":
        message = request.POST.get("message", "").strip()
        if not message:
            return JsonResponse({"error": "Empty message"}, status=400)

        # Check if user has processed documents
        has_documents = documents.filter(is_processed=True).exists()

        if has_documents:
            # Search in user's documents only
            context_chunks = search_documents(request.user.id, message)
            if context_chunks:
                print(f"Found {len(context_chunks)} relevant chunks for query")
            response = ask_ollama(message, context_chunks)
        else:
            # Regular chat without document context
            response = ask_ollama(message)

        # Save chat for this user
        Chat.objects.create(
            user=request.user,
            message=message,
            response=response,
            created_at=timezone.now()
        )

        return JsonResponse({"message": message, "response": response})

    return render(request, "chatbot.html", {"chats": chats, "documents": documents})


# ================================
# Document Upload
# ================================
def upload_document(request):
    """Handle document upload for specific user."""
    if request.method == "POST" and request.FILES.get("document"):
        file = request.FILES["document"]
        allowed_extensions = [".pdf", ".docx", ".txt"]
        file_ext = os.path.splitext(file.name)[1].lower()

        if file_ext not in allowed_extensions:
            return JsonResponse({
                "error": "Invalid file type. Only PDF, DOCX, and TXT files are allowed."
            }, status=400)

        # Check file size (limit to 10MB)
        if file.size > 10 * 1024 * 1024:
            return JsonResponse({
                "error": "File too large. Maximum size is 10MB."
            }, status=400)

        # Create document record for this user
        document = Document.objects.create(
            user=request.user,
            title=file.name,
            file=file
        )

        # Process document
        success = process_document(document.id, request.user.id)
        
        if success:
            return JsonResponse({
                "success": True,
                "message": "Document uploaded and processed successfully!",
                "document_id": document.id
            })
        else:
            # Clean up if processing failed
            if document.file:
                document.file.delete()
            document.delete()
            return JsonResponse({
                "error": "Failed to process document. Please ensure it contains readable text."
            }, status=400)

    return JsonResponse({"error": "No file provided"}, status=400)


# ================================
# Document Delete
# ================================
def delete_document(request, document_id):
    """Delete a document and its embeddings for specific user."""
    try:
        # Ensure document belongs to the requesting user
        document = Document.objects.get(id=document_id, user=request.user)
        
        # Delete from vector database
        delete_user_document_chunks(request.user.id, document_id)

        # Delete file from storage
        if document.file:
            document.file.delete()
            
        # Delete database record
        document.delete()

        return JsonResponse({
            "success": True,
            "message": "Document deleted successfully"
        })
        
    except Document.DoesNotExist:
        return JsonResponse({
            "error": "Document not found or you don't have permission to delete it."
        }, status=404)
    except Exception as e:
        print(f"Error deleting document: {e}")
        return JsonResponse({"error": "Failed to delete document."}, status=400)


# ================================
# Auth Views
# ================================
def login(request):
    if request.method == "POST":
        username = request.POST["username"]
        password = request.POST["password"]
        user = auth.authenticate(request, username=username, password=password)
        if user:
            auth.login(request, user)
            return redirect("chatbot")
        return render(request, "login.html", {"error_message": "Invalid username or password"})
    return render(request, "login.html")


def register(request):
    """Handle new user registration."""
    if request.method == "POST":
        username = request.POST.get("username", "").strip()
        email = request.POST.get("email", "").strip()
        password1 = request.POST.get("password1", "")
        password2 = request.POST.get("password2", "")

        # Validation
        if not username or not email or not password1 or not password2:
            messages.error(request, "All fields are required.")
        elif password1 != password2:
            messages.error(request, "Passwords do not match.")
        elif User.objects.filter(username=username).exists():
            messages.error(request, "Username already taken.")
        elif User.objects.filter(email=email).exists():
            messages.error(request, "Email already in use.")
        else:
            user = User.objects.create_user(username=username, email=email, password=password1)
            auth.login(request, user)
            return redirect("chatbot")

    return render(request, "register.html")



def logout(request):
    auth.logout(request)
    return redirect("login")